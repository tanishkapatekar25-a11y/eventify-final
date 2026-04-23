from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_PATH = "Eventify-Project-Report.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_title(document):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("EVENTIFY\nProject Report")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x11, 0x8A, 0x7E)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "A Readymade Event Information Website Developed Using HTML, CSS, and JavaScript"
    )
    run.italic = True
    run.font.size = Pt(13)

    document.add_paragraph("")
    info = document.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run("Prepared for project presentation and documentation\n").bold = True
    info.add_run("Project Type: Frontend Static Website")

    document.add_page_break()


def add_heading(document, text, level=1):
    heading = document.add_heading(level=level)
    run = heading.add_run(text)
    run.font.color.rgb = RGBColor(0x11, 0x8A, 0x7E)


def add_bullets(document, items):
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def build_report():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)

    add_title(doc)

    add_heading(doc, "1. Introduction")
    doc.add_paragraph(
        "Eventify is a modern multi-page event information website created to help users "
        "discover events, browse by category, check schedules, and complete registrations "
        "through a clean and professional interface. The project was designed as a static "
        "frontend website that demonstrates strong layout planning, responsive design, "
        "interactive features, and scalable structure."
    )
    doc.add_paragraph(
        "The website focuses on five major categories of events: Sports, Music, Education, "
        "Travel, and Seminar. Each category contains three events, producing a total of "
        "fifteen event pages with detailed information and registration forms."
    )

    add_heading(doc, "2. Objectives of the Project")
    add_bullets(
        doc,
        [
            "To design a professional event information website with a clean user interface.",
            "To provide category-wise event browsing and dedicated event detail pages.",
            "To add interactive frontend features such as countdown timers, forms, search, filters, and popups.",
            "To demonstrate responsive design for desktop, tablet, and mobile screens.",
            "To build a project structure that is easy to maintain and scale in the future.",
        ],
    )

    add_heading(doc, "3. Project Overview")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Project Item"
    hdr[1].text = "Details"
    for cell in hdr:
        set_cell_shading(cell, "D9F2EE")

    rows = [
        ("Project Name", "Eventify"),
        ("Project Type", "Static Event Information Website"),
        ("Main Technologies", "HTML, CSS, JavaScript"),
        ("Total Event Categories", "5"),
        ("Total Event Pages", "15"),
        ("Additional Pages", "About, Contact, FAQ, Testimonials, Dashboard, Documentation"),
        ("Theme", "Teal-based professional visual theme"),
        ("Special Features", "Dark/Light mode, theme switcher, search, filters, countdowns, popups"),
    ]
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value

    add_heading(doc, "4. Website Structure")
    doc.add_paragraph(
        "The Eventify website is organized into shared assets and multiple pages. Shared "
        "files provide styling, interactivity, and event data, while individual pages "
        "present the actual content to users."
    )
    add_bullets(
        doc,
        [
            "Home Page: hero banner, slideshow, navigation, quick links, and upcoming event countdown.",
            "Categories Page: search and filter tools with grouped event cards.",
            "15 Event Detail Pages: one page per event with description, countdown, and registration form.",
            "About Page: purpose, organizers, and mission statement.",
            "Contact Page: contact form and embedded venue map.",
            "FAQ Page: collapsible frequently asked questions.",
            "Testimonials Page: feedback from past participants.",
            "Dashboard Page: summary metrics and ticket statistics.",
            "Documentation Page: explanation of project structure and technologies used.",
        ],
    )

    add_heading(doc, "5. Event Categories Included")
    cat_table = doc.add_table(rows=1, cols=2)
    cat_table.style = "Table Grid"
    cat_hdr = cat_table.rows[0].cells
    cat_hdr[0].text = "Category"
    cat_hdr[1].text = "Events"
    for cell in cat_hdr:
        set_cell_shading(cell, "D9F2EE")

    categories = [
        ("Sports", "Intercollege Football Tournament, Annual Marathon, Cricket League"),
        ("Music", "Rock Concert, Classical Night, DJ Party"),
        ("Education", "Science Fair, Coding Hackathon, Math Olympiad"),
        ("Travel", "Adventure Trek, City Heritage Walk, Beach Camping"),
        ("Seminar", "Entrepreneurship Seminar, Data Science Workshop, Leadership Talk"),
    ]
    for category, events in categories:
        cells = cat_table.add_row().cells
        cells[0].text = category
        cells[1].text = events

    add_heading(doc, "6. Major Features Implemented")
    add_heading(doc, "6.1 Visual and Design Features", level=2)
    add_bullets(
        doc,
        [
            "Teal color theme for brand consistency.",
            "Professional card-based layout suitable for academic presentation.",
            "Hover effects on event cards for improved visual feedback.",
            "Responsive design for desktop, tablet, and mobile devices.",
            "Theme controls for dark/light mode and teal/neutral palette selection.",
        ],
    )

    add_heading(doc, "6.2 Interactive Features", level=2)
    add_bullets(
        doc,
        [
            "Homepage slideshow carousel for event highlights.",
            "Countdown timer for the next upcoming event on the homepage.",
            "Per-event countdown timers on all dedicated event pages.",
            "Search and filter functionality by keyword, category, and month.",
            "Registration form with popup-based ticket confirmation.",
            "Ticket badge feedback such as Early Bird Registered and VIP Ticket Holder.",
            "FAQ collapsible sections for smoother navigation.",
        ],
    )

    add_heading(doc, "6.3 Professional and Extra Features", level=2)
    add_bullets(
        doc,
        [
            "Embedded Google Map on the contact page.",
            "Newsletter signup section on the homepage.",
            "Dashboard page for event statistics.",
            "Documentation page describing structure, features, and scalability.",
            "Social media links for added professionalism.",
            "Readable typography, spacing, and focus states to support accessibility.",
        ],
    )

    add_heading(doc, "7. Technologies Used")
    doc.add_paragraph(
        "The project was developed using core web technologies without relying on a backend "
        "framework. This keeps the website easy to understand and suitable for frontend "
        "project evaluation."
    )
    tech_table = doc.add_table(rows=1, cols=2)
    tech_table.style = "Table Grid"
    tech_hdr = tech_table.rows[0].cells
    tech_hdr[0].text = "Technology"
    tech_hdr[1].text = "Purpose"
    for cell in tech_hdr:
        set_cell_shading(cell, "D9F2EE")

    tech_rows = [
        ("HTML", "Used to create page structure and semantic content."),
        ("CSS", "Used for layout, color themes, responsiveness, animations, and styling."),
        ("JavaScript", "Used for theme switching, slideshow, countdowns, filters, forms, and dashboard logic."),
    ]
    for tech, purpose in tech_rows:
        cells = tech_table.add_row().cells
        cells[0].text = tech
        cells[1].text = purpose

    add_heading(doc, "8. Responsive Design and Accessibility")
    doc.add_paragraph(
        "Responsive design was an important part of the Eventify website. Flexible layouts, "
        "collapsible navigation, adaptive grids, and readable spacing allow the website to "
        "work properly on different screen sizes. Accessibility was improved through readable "
        "font choices, visible focus states, sufficient spacing, and color contrast designed "
        "for both light and dark modes."
    )

    add_heading(doc, "9. Scalability and Maintainability")
    doc.add_paragraph(
        "The project was structured to make future updates easy. The stylesheet is shared "
        "across all pages, the main JavaScript file handles common interactions, and the "
        "event dataset is centralized in a reusable data file. Because of this structure, "
        "new events can be added by extending the shared event data and following the same "
        "event page pattern."
    )
    add_bullets(
        doc,
        [
            "Shared stylesheet for consistent layout and branding.",
            "Shared JavaScript for common user interactions.",
            "Central event data for rendering and dashboard summaries.",
            "Reusable event page template for similar page creation.",
        ],
    )

    add_heading(doc, "10. Output and Benefits")
    add_paragraph = doc.add_paragraph
    add_paragraph(
        "The final website demonstrates not only frontend development skills but also project "
        "organization, user experience planning, design consistency, and professional presentation. "
        "It is suitable for academic submission, portfolio demonstration, and further expansion."
    )
    add_bullets(
        doc,
        [
            "Provides a complete multi-page website instead of a single static page.",
            "Improves user convenience through filters, countdowns, and registration flow.",
            "Maintains a polished and examiner-friendly interface.",
            "Shows practical understanding of HTML, CSS, and JavaScript integration.",
        ],
    )

    add_heading(doc, "11. Conclusion")
    doc.add_paragraph(
        "Eventify is a complete readymade event information website that combines design, "
        "interactivity, and structured content management in one frontend project. The "
        "website includes five event categories, fifteen event pages, multiple utility pages, "
        "theme controls, filters, countdown timers, and registration popups. Overall, the "
        "project successfully meets the objective of building a modern, scalable, and "
        "presentation-ready event website using HTML, CSS, and JavaScript."
    )

    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_report()
